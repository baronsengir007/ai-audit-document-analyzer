"""
Document Processor Module

Handles document loading, text extraction, and content normalization 
for various document types (PDF, DOCX, XLSX).
"""

import logging
from pathlib import Path
from typing import Dict, Any, Optional, Union, List

import PyPDF2
from docx import Document as DocxDocument
import openpyxl

from .interfaces import Document, DocumentProcessor as IDocumentProcessor


class DocumentProcessor(IDocumentProcessor):
    """
    Handles document processing including text extraction and normalization
    for various document types
    """
    
    def __init__(self, logger: Optional[logging.Logger] = None):
        """
        Initialize DocumentProcessor
        
        Args:
            logger: Optional logger instance
        """
        self.logger = logger or logging.getLogger(__name__)
    
    def process_document(self, doc_path: Path) -> Document:
        """
        Process a document and return a unified Document object
        
        Args:
            doc_path: Path to the document
            
        Returns:
            Document object with extracted content
        """
        try:
            filename = doc_path.name
            content = self.extract_text(doc_path)
            
            # Create and return Document object
            return Document(
                filename=filename,
                content=content,
                classification="",  # To be determined later
                metadata=self._extract_metadata(doc_path),
                source_path=str(doc_path)
            )
        except Exception as e:
            self.logger.error(f"Error processing document {doc_path}: {e}")
            # Return a Document with error information
            return Document(
                filename=doc_path.name,
                content=f"[PROCESSING ERROR] {str(e)}",
                classification="error",
                metadata={"error": str(e)},
                source_path=str(doc_path)
            )
    
    def extract_text(self, file_path: Union[str, Path]) -> str:
        """
        Extract text from a document based on its file extension
        
        Args:
            file_path: Path to the document
            
        Returns:
            Extracted text content
        """
        file_path = Path(file_path) if isinstance(file_path, str) else file_path
        
        if not file_path.exists():
            self.logger.error(f"File not found: {file_path}")
            return "[FILE NOT FOUND]"
        
        ext = file_path.suffix.lower()
        
        try:
            if ext == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif ext == '.docx':
                return self.extract_text_from_word(file_path)
            elif ext == '.xlsx':
                return self.extract_text_from_excel(file_path)
            else:
                self.logger.warning(f"Unsupported file extension: {ext}")
                return f"[UNSUPPORTED FILE TYPE: {ext}]"
        except Exception as e:
            self.logger.error(f"Error extracting text from {file_path}: {e}")
            return f"[EXTRACTION ERROR] {str(e)}"
    
    def extract_text_from_pdf(self, pdf_path: Union[str, Path]) -> str:
        """
        Extract text from a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            with open(pdf_path, "rb") as file:
                try:
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        extracted_text = page.extract_text()
                        if extracted_text:
                            text += extracted_text + "\n"
                except (PyPDF2.errors.PdfReadError, Exception) as e:
                    self.logger.warning(f"Error reading PDF content from {pdf_path}: {e}")
                    # Attempt a more basic extraction method if possible
                    file.seek(0)  # Reset file pointer
                    try:
                        # Try a more basic extraction to at least get some content
                        raw_data = file.read().decode('utf-8', errors='ignore')
                        # Return first 2000 characters as a fallback
                        text = f"[PDF EXTRACTION PARTIAL] {raw_data[:2000]}"
                    except Exception as fallback_error:
                        self.logger.error(f"Fallback extraction failed: {fallback_error}")
                        # Return placeholder text with error info
                        text = f"[PDF EXTRACTION FAILED] This document could not be processed properly. Error: {str(e)}"
        except FileNotFoundError:
            self.logger.error(f"PDF file not found: {pdf_path}")
            text = "[PDF FILE NOT FOUND]"
        except Exception as e:
            self.logger.error(f"Unexpected error opening {pdf_path}: {e}")
            text = f"[PDF PROCESSING ERROR] {str(e)}"
            
        # Ensure we always return some text content
        if not text:
            text = "[EMPTY PDF CONTENT]"
            
        return text
    
    def extract_text_from_word(self, docx_path: Union[str, Path]) -> str:
        """
        Extract text from a Word document
        
        Args:
            docx_path: Path to the Word document
            
        Returns:
            Extracted text content
        """
        text = ""
        try:
            doc = DocxDocument(docx_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text]
            text = "\n".join(paragraphs)
            
            # If document has tables, extract table content
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text += "\n" + " | ".join(row_text)
            
        except FileNotFoundError:
            self.logger.error(f"Word file not found: {docx_path}")
            text = "[WORD FILE NOT FOUND]"
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document {docx_path}: {e}")
            text = f"[WORD PROCESSING ERROR] {str(e)}"
        
        # Ensure we always return some text content
        if not text:
            text = "[EMPTY WORD CONTENT]"
            
        return text
    
    def extract_text_from_excel(self, xlsx_path: Union[str, Path]) -> str:
        """
        Extract text from an Excel document
        
        Args:
            xlsx_path: Path to the Excel document
            
        Returns:
            Extracted text content
        """
        text_lines = []
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_content = []
                
                # Add sheet name as header
                sheet_content.append(f"--- Sheet: {sheet_name} ---")
                
                # Process each row
                for row in ws.iter_rows(values_only=True):
                    row_text = [str(cell) for cell in row if cell is not None]
                    if row_text:
                        sheet_content.append(" | ".join(row_text))
                
                # Add sheet content if not empty
                if len(sheet_content) > 1:  # More than just the header
                    text_lines.extend(sheet_content)
                    text_lines.append("")  # Empty line between sheets
            
        except FileNotFoundError:
            self.logger.error(f"Excel file not found: {xlsx_path}")
            text_lines = ["[EXCEL FILE NOT FOUND]"]
        except Exception as e:
            self.logger.error(f"Error extracting text from Excel document {xlsx_path}: {e}")
            text_lines = [f"[EXCEL PROCESSING ERROR] {str(e)}"]
        
        # Ensure we always return some text content
        if not text_lines:
            text_lines = ["[EMPTY EXCEL CONTENT]"]
            
        return "\n".join(text_lines)
    
    def _extract_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from a document
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary of metadata
        """
        metadata = {
            "file_size": self._get_file_size(file_path),
            "file_extension": file_path.suffix.lower(),
            "last_modified": self._get_last_modified(file_path)
        }
        
        # Extract additional metadata based on file type
        try:
            if file_path.suffix.lower() == '.pdf':
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif file_path.suffix.lower() == '.docx':
                docx_metadata = self._extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
            elif file_path.suffix.lower() == '.xlsx':
                xlsx_metadata = self._extract_xlsx_metadata(file_path)
                metadata.update(xlsx_metadata)
        except Exception as e:
            self.logger.warning(f"Error extracting metadata from {file_path}: {e}")
            metadata["metadata_error"] = str(e)
        
        return metadata
    
    def _get_file_size(self, file_path: Path) -> int:
        """Get file size in bytes"""
        try:
            return file_path.stat().st_size
        except Exception:
            return 0
    
    def _get_last_modified(self, file_path: Path) -> str:
        """Get last modified timestamp as string"""
        try:
            return file_path.stat().st_mtime
        except Exception:
            return ""
    
    def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract PDF-specific metadata"""
        metadata = {}
        try:
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                info = reader.metadata
                if info:
                    for key, value in info.items():
                        if value and isinstance(value, (str, int, float, bool)):
                            metadata[key[1:] if key.startswith('/') else key] = value
                
                metadata["page_count"] = len(reader.pages)
        except Exception as e:
            self.logger.warning(f"Error extracting PDF metadata: {e}")
            metadata["metadata_extraction_error"] = str(e)
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract DOCX-specific metadata"""
        metadata = {}
        try:
            doc = DocxDocument(file_path)
            props = doc.core_properties
            
            # Extract common properties
            if props.author:
                metadata["author"] = props.author
            if props.title:
                metadata["title"] = props.title
            if props.created:
                metadata["created"] = str(props.created)
            if props.modified:
                metadata["modified"] = str(props.modified)
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
            
            # Count paragraphs and tables
            metadata["paragraph_count"] = len(doc.paragraphs)
            metadata["table_count"] = len(doc.tables)
        except Exception as e:
            self.logger.warning(f"Error extracting DOCX metadata: {e}")
            metadata["metadata_extraction_error"] = str(e)
        
        return metadata
    
    def _extract_xlsx_metadata(self, file_path: Path) -> Dict[str, Any]:
        """Extract XLSX-specific metadata"""
        metadata = {}
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            # Basic workbook info
            metadata["sheet_count"] = len(wb.sheetnames)
            metadata["sheet_names"] = wb.sheetnames
            
            # Count cells with content
            total_cells = 0
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                sheet_cells = sum(1 for row in ws.iter_rows() for cell in row if cell.value is not None)
                metadata[f"cells_in_{sheet_name}"] = sheet_cells
                total_cells += sheet_cells
            
            metadata["total_cells"] = total_cells
        except Exception as e:
            self.logger.warning(f"Error extracting XLSX metadata: {e}")
            metadata["metadata_extraction_error"] = str(e)
        
        return metadata


# For backward compatibility, keep the original function interfaces
def extract_text_from_pdf(pdf_path):
    """Extract text from PDF (for backward compatibility)"""
    processor = DocumentProcessor()
    return processor.extract_text_from_pdf(pdf_path)

def extract_text_from_word(docx_path):
    """Extract text from Word (for backward compatibility)"""
    processor = DocumentProcessor()
    return processor.extract_text_from_word(docx_path)

def extract_text_from_excel(xlsx_path):
    """Extract text from Excel (for backward compatibility)"""
    processor = DocumentProcessor()
    return processor.extract_text_from_excel(xlsx_path)


# For testing purpose
if __name__ == "__main__":
    processor = DocumentProcessor()
    
    # Test PDF extraction
    pdf_path = Path("docs/KPN FEB 2025.pdf")
    if pdf_path.exists():
        text_pdf = processor.extract_text_from_pdf(str(pdf_path))
        print("\n--- PDF Document Text ---\n")
        print(text_pdf)
    else:
        print(f"PDF file not found: {pdf_path}")

    # Test Word extraction
    word_path = Path("docs/Request for Information Alpha.docx")
    if word_path.exists():
        text_word = processor.extract_text_from_word(str(word_path))
        print("\n--- Word Document Text ---\n")
        print(text_word)
    else:
        print(f"Word file not found: {word_path}")

    # Test Excel extraction
    excel_path = Path("docs/pm-Project Alpha.xlsx")
    if excel_path.exists():
        text_excel = processor.extract_text_from_excel(str(excel_path))
        print("\n--- Excel Document Text ---\n")
        print(text_excel)
    else:
        print(f"Excel file not found: {excel_path}")