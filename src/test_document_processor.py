"""
Unit tests for document_processor.py
"""

import unittest
import tempfile
import os
from pathlib import Path
from unittest.mock import patch, MagicMock, mock_open

import PyPDF2
from docx import Document as DocxDocument

from .document_processor import DocumentProcessor
from .interfaces import Document


class TestDocumentProcessor(unittest.TestCase):
    """Test the DocumentProcessor class"""
    
    def setUp(self):
        """Set up test fixtures"""
        self.processor = DocumentProcessor()
        
        # Create temp directory for test files
        self.temp_dir = tempfile.TemporaryDirectory()
        self.test_dir = Path(self.temp_dir.name)
        
        # Create test files
        self.pdf_path = self.test_dir / "test.pdf"
        self.docx_path = self.test_dir / "test.docx"
        self.xlsx_path = self.test_dir / "test.xlsx"
        self.unknown_path = self.test_dir / "test.unknown"
        
        # Touch files to create them
        self.pdf_path.touch()
        self.docx_path.touch()
        self.xlsx_path.touch()
        self.unknown_path.touch()
    
    def tearDown(self):
        """Clean up after tests"""
        self.temp_dir.cleanup()
    
    def test_process_document(self):
        """Test process_document method"""
        with patch.object(self.processor, 'extract_text', return_value="Test content"):
            with patch.object(self.processor, '_extract_metadata', return_value={"test": "metadata"}):
                doc = self.processor.process_document(self.pdf_path)
                
                self.assertIsInstance(doc, Document)
                self.assertEqual(doc.filename, "test.pdf")
                self.assertEqual(doc.content, "Test content")
                self.assertEqual(doc.metadata, {"test": "metadata"})
                self.assertEqual(doc.source_path, str(self.pdf_path))
    
    def test_process_document_error(self):
        """Test process_document error handling"""
        with patch.object(self.processor, 'extract_text', side_effect=Exception("Test error")):
            doc = self.processor.process_document(self.pdf_path)
            
            self.assertIsInstance(doc, Document)
            self.assertEqual(doc.classification, "error")
            self.assertIn("[PROCESSING ERROR]", doc.content)
            self.assertIn("Test error", doc.content)
    
    def test_extract_text_pdf(self):
        """Test extract_text for PDF files"""
        with patch.object(self.processor, 'extract_text_from_pdf', return_value="PDF content"):
            content = self.processor.extract_text(self.pdf_path)
            self.assertEqual(content, "PDF content")
    
    def test_extract_text_docx(self):
        """Test extract_text for DOCX files"""
        with patch.object(self.processor, 'extract_text_from_word', return_value="DOCX content"):
            content = self.processor.extract_text(self.docx_path)
            self.assertEqual(content, "DOCX content")
    
    def test_extract_text_xlsx(self):
        """Test extract_text for XLSX files"""
        with patch.object(self.processor, 'extract_text_from_excel', return_value="XLSX content"):
            content = self.processor.extract_text(self.xlsx_path)
            self.assertEqual(content, "XLSX content")
    
    def test_extract_text_unknown(self):
        """Test extract_text for unknown file types"""
        content = self.processor.extract_text(self.unknown_path)
        self.assertIn("[UNSUPPORTED FILE TYPE", content)
    
    def test_extract_text_missing_file(self):
        """Test extract_text for missing files"""
        missing_path = self.test_dir / "missing.pdf"
        content = self.processor.extract_text(missing_path)
        self.assertIn("[FILE NOT FOUND]", content)
    
    def test_extract_text_from_pdf(self):
        """Test extract_text_from_pdf"""
        # Mock the PDF reader
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page content"
        
        mock_reader = MagicMock()
        mock_reader.pages = [mock_page, mock_page]  # Two pages
        
        with patch('PyPDF2.PdfReader', return_value=mock_reader):
            content = self.processor.extract_text_from_pdf(self.pdf_path)
            self.assertEqual(content, "Page content\nPage content\n")
    
    def test_extract_text_from_pdf_error(self):
        """Test extract_text_from_pdf error handling"""
        # Mock PDF reader to raise an exception
        with patch('PyPDF2.PdfReader', side_effect=PyPDF2.errors.PdfReadError("Test error")):
            with patch('builtins.open', mock_open(read_data=b"Some PDF data")):
                content = self.processor.extract_text_from_pdf(self.pdf_path)
                self.assertIn("[PDF EXTRACTION PARTIAL]", content)
    
    def test_extract_text_from_word(self):
        """Test extract_text_from_word"""
        # Mock the Word document
        mock_para1 = MagicMock()
        mock_para1.text = "Paragraph 1"
        
        mock_para2 = MagicMock()
        mock_para2.text = "Paragraph 2"
        
        mock_cell1 = MagicMock()
        mock_cell1.text = "Cell 1"
        
        mock_cell2 = MagicMock()
        mock_cell2.text = "Cell 2"
        
        mock_row = MagicMock()
        mock_row.cells = [mock_cell1, mock_cell2]
        
        mock_table = MagicMock()
        mock_table.rows = [mock_row]
        
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_doc.tables = [mock_table]
        
        with patch('docx.Document', return_value=mock_doc):
            content = self.processor.extract_text_from_word(self.docx_path)
            self.assertEqual(content, "Paragraph 1\nParagraph 2\nCell 1 | Cell 2")
    
    def test_extract_text_from_word_error(self):
        """Test extract_text_from_word error handling"""
        with patch('docx.Document', side_effect=Exception("Test error")):
            content = self.processor.extract_text_from_word(self.docx_path)
            self.assertIn("[WORD PROCESSING ERROR]", content)
    
    def test_extract_text_from_excel(self):
        """Test extract_text_from_excel"""
        # Mock the Excel workbook
        mock_ws = MagicMock()
        mock_ws.iter_rows.return_value = [
            ("Cell A1", "Cell B1"),
            ("Cell A2", "Cell B2")
        ]
        
        mock_wb = MagicMock()
        mock_wb.sheetnames = ["Sheet1", "Sheet2"]
        mock_wb.__getitem__.return_value = mock_ws
        
        with patch('openpyxl.load_workbook', return_value=mock_wb):
            content = self.processor.extract_text_from_excel(self.xlsx_path)
            self.assertIn("Sheet: Sheet1", content)
            self.assertIn("Cell A1 | Cell B1", content)
            self.assertIn("Cell A2 | Cell B2", content)
    
    def test_extract_text_from_excel_error(self):
        """Test extract_text_from_excel error handling"""
        with patch('openpyxl.load_workbook', side_effect=Exception("Test error")):
            content = self.processor.extract_text_from_excel(self.xlsx_path)
            self.assertIn("[EXCEL PROCESSING ERROR]", content)
    
    def test_extract_metadata(self):
        """Test _extract_metadata"""
        with patch.object(self.processor, '_get_file_size', return_value=1024):
            with patch.object(self.processor, '_get_last_modified', return_value=12345):
                with patch.object(self.processor, '_extract_pdf_metadata', return_value={"pdf": "metadata"}):
                    metadata = self.processor._extract_metadata(self.pdf_path)
                    self.assertEqual(metadata["file_size"], 1024)
                    self.assertEqual(metadata["file_extension"], ".pdf")
                    self.assertEqual(metadata["last_modified"], 12345)
                    self.assertEqual(metadata["pdf"], "metadata")
    
    def test_extract_pdf_metadata(self):
        """Test _extract_pdf_metadata"""
        # Mock PDF metadata
        mock_info = {"/Title": "Test Title", "/Author": "Test Author"}
        
        mock_reader = MagicMock()
        mock_reader.metadata = mock_info
        mock_reader.pages = [1, 2, 3]  # 3 pages
        
        with patch('PyPDF2.PdfReader', return_value=mock_reader):
            metadata = self.processor._extract_pdf_metadata(self.pdf_path)
            self.assertEqual(metadata["Title"], "Test Title")
            self.assertEqual(metadata["Author"], "Test Author")
            self.assertEqual(metadata["page_count"], 3)
    
    def test_extract_docx_metadata(self):
        """Test _extract_docx_metadata"""
        # Mock DOCX metadata
        mock_props = MagicMock()
        mock_props.author = "Test Author"
        mock_props.title = "Test Title"
        mock_props.created = "2023-01-01"
        mock_props.modified = "2023-01-02"
        mock_props.subject = "Test Subject"
        mock_props.keywords = "test, keywords"
        
        mock_doc = MagicMock()
        mock_doc.core_properties = mock_props
        mock_doc.paragraphs = [1, 2, 3]  # 3 paragraphs
        mock_doc.tables = [1]  # 1 table
        
        with patch('docx.Document', return_value=mock_doc):
            metadata = self.processor._extract_docx_metadata(self.docx_path)
            self.assertEqual(metadata["author"], "Test Author")
            self.assertEqual(metadata["title"], "Test Title")
            self.assertEqual(metadata["created"], "2023-01-01")
            self.assertEqual(metadata["paragraph_count"], 3)
            self.assertEqual(metadata["table_count"], 1)
    
    def test_extract_xlsx_metadata(self):
        """Test _extract_xlsx_metadata"""
        # Mock XLSX metadata
        mock_ws = MagicMock()
        mock_ws.iter_rows.return_value = [
            [MagicMock(value="A"), MagicMock(value="B")],
            [MagicMock(value="C"), MagicMock(value=None)]
        ]
        
        mock_wb = MagicMock()
        mock_wb.sheetnames = ["Sheet1", "Sheet2"]
        mock_wb.__getitem__.return_value = mock_ws
        
        with patch('openpyxl.load_workbook', return_value=mock_wb):
            metadata = self.processor._extract_xlsx_metadata(self.xlsx_path)
            self.assertEqual(metadata["sheet_count"], 2)
            self.assertEqual(metadata["sheet_names"], ["Sheet1", "Sheet2"])
            self.assertEqual(metadata["cells_in_Sheet1"], 3)  # 3 non-None cells


if __name__ == "__main__":
    unittest.main()