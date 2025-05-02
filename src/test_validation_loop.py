import unittest
from pathlib import Path
import tempfile
import shutil
import json
import yaml
import time
import sys
from pathlib import Path
import PyPDF2
from docx import Document
import openpyxl
import io
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Add the src directory to the Python path
src_path = str(Path(__file__).parent)
if src_path not in sys.path:
    sys.path.append(src_path)

from validation_loop import ValidationLoop, ValidationConfig

class TestValidationLoop(unittest.TestCase):
    def setUp(self):
        """Set up test environment"""
        # Create temporary directory for test files
        self.test_dir = tempfile.mkdtemp()
        self.config = ValidationConfig(
            strict_mode=False,
            batch_size=2,
            max_workers=2,
            log_level="DEBUG",
            output_format="json",
            detailed_report=True
        )
        
        # Create test checklist
        self.checklist_path = Path(self.test_dir) / "config"
        self.checklist_path.mkdir(exist_ok=True)
        self.checklist_file = self.checklist_path / "checklist.yaml"
        self.checklist_data = {
            "audit_completeness_checklist": [
                {
                    "id": "invoices",
                    "required_keywords": ["invoice", "total", "date"]
                },
                {
                    "id": "audit_rfi",
                    "required_keywords": ["audit", "requirement", "compliance"]
                },
                {
                    "id": "project_data",
                    "required_keywords": ["project", "status", "milestone"]
                }
            ]
        }
        with open(self.checklist_file, "w") as f:
            yaml.dump(self.checklist_data, f)

        # Create test documents
        self.doc_dir = Path(self.test_dir) / "docs"
        self.doc_dir.mkdir()

        # Create test PDF
        self.pdf_path = self.doc_dir / "test_invoice.pdf"
        c = canvas.Canvas(str(self.pdf_path), pagesize=letter)
        c.drawString(100, 750, "This is an invoice document")
        c.drawString(100, 700, "Total amount: $1000")
        c.drawString(100, 650, "Date: 2024-01-01")
        c.save()

        # Create test Word document
        self.word_path = self.doc_dir / "test_audit.docx"
        doc = Document()
        doc.add_paragraph("Audit Requirements Document")
        doc.add_paragraph("Compliance checklist")
        doc.add_paragraph("Requirement 1: Security")
        doc.save(self.word_path)

        # Create test Excel document
        self.excel_path = self.doc_dir / "test_project.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Project Status Report"
        ws["A2"] = "Milestone 1: Completed"
        ws["A3"] = "Status: On track"
        wb.save(self.excel_path)

        # Create empty document
        self.empty_path = self.doc_dir / "empty.pdf"
        c = canvas.Canvas(str(self.empty_path), pagesize=letter)
        c.save()

        # Create invalid document
        self.invalid_path = self.doc_dir / "invalid.txt"
        with open(self.invalid_path, "w") as f:
            f.write("This is not a supported document type")

    def tearDown(self):
        """Clean up test environment"""
        shutil.rmtree(self.test_dir)

    def test_validation_loop_initialization(self):
        """Test validation loop initialization"""
        validator = ValidationLoop(self.config)
        self.assertIsNotNone(validator.checklist_map)
        self.assertIsNotNone(validator.type_to_checklist_id)

    def test_process_single_document(self):
        """Test processing a single document"""
        validator = ValidationLoop(self.config)
        result = validator.process_document(self.pdf_path)
        
        self.assertEqual(result["status"], "complete")
        self.assertEqual(result["document"]["filename"], "test_invoice.pdf")
        self.assertTrue("validation_results" in result)

    def test_process_batch(self):
        """Test processing a batch of documents"""
        validator = ValidationLoop(self.config)
        results = validator.process_batch([self.pdf_path, self.word_path])
        
        self.assertEqual(len(results), 2)
        self.assertTrue(all(r["status"] in ["complete", "incomplete"] for r in results))

    def test_validate_documents(self):
        """Test full document validation process"""
        validator = ValidationLoop(self.config)
        results = validator.validate_documents(self.doc_dir)
        
        self.assertEqual(results["status"], "success")
        self.assertEqual(results["summary"]["total_documents"], 5)  # Including empty and invalid
        self.assertTrue(results["summary"]["complete_documents"] >= 1)
        self.assertTrue(results["summary"]["error_documents"] >= 1)  # At least invalid.txt

    def test_empty_directory(self):
        """Test validation with empty directory"""
        empty_dir = Path(self.test_dir) / "empty"
        empty_dir.mkdir()
        
        validator = ValidationLoop(self.config)
        results = validator.validate_documents(empty_dir)
        
        self.assertEqual(results["status"], "no_documents")
        self.assertEqual(len(results["results"]), 0)

    def test_invalid_directory(self):
        """Test validation with non-existent directory"""
        validator = ValidationLoop(self.config)
        with self.assertRaises(ValueError):
            validator.validate_documents(Path(self.test_dir) / "nonexistent")

    def test_save_results(self):
        """Test saving validation results"""
        validator = ValidationLoop(self.config)
        results = validator.validate_documents(self.doc_dir)
        
        output_path = Path(self.test_dir) / "results.json"
        validator.save_results(results, output_path)
        
        self.assertTrue(output_path.exists())
        with open(output_path) as f:
            saved_results = json.load(f)
            self.assertEqual(saved_results["status"], "success")

    def test_strict_mode(self):
        """Test validation in strict mode"""
        strict_config = ValidationConfig(strict_mode=True)
        validator = ValidationLoop(strict_config)
        result = validator.process_document(self.pdf_path)
        
        # In strict mode, even minor issues should be marked as incomplete
        self.assertEqual(result["status"], "incomplete")

    def test_parallel_processing(self):
        """Test parallel document processing"""
        validator = ValidationLoop(self.config)
        start_time = time.time()
        results = validator.validate_documents(self.doc_dir)
        end_time = time.time()
        
        # Verify all documents were processed
        self.assertEqual(len(results["results"]), 5)
        # Verify parallel processing was used (should be faster than sequential)
        self.assertLess(end_time - start_time, 5.0)  # Adjust threshold as needed

    def test_error_handling(self):
        """Test error handling for various scenarios"""
        validator = ValidationLoop(self.config)
        
        # Test with empty document
        result = validator.process_document(self.empty_path)
        self.assertEqual(result["status"], "error")
        
        # Test with invalid document type
        result = validator.process_document(self.invalid_path)
        self.assertEqual(result["status"], "error")

    def test_document_classification(self):
        """Test document classification accuracy"""
        validator = ValidationLoop(self.config)
        
        # Test invoice classification
        result = validator.process_document(self.pdf_path)
        self.assertEqual(result["document"]["classification"], "invoice")
        
        # Test audit document classification
        result = validator.process_document(self.word_path)
        self.assertEqual(result["document"]["classification"], "audit_rfi")
        
        # Test project data classification
        result = validator.process_document(self.excel_path)
        self.assertEqual(result["document"]["classification"], "project_data")

if __name__ == "__main__":
    unittest.main() 